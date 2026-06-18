from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

out_path = 'FinAgentPro交互设计与流程图.docx'
img_dir = Path('diagrams/2026-06-18T143000')

doc = Document()

# Title
title = doc.add_heading('FinAgent Pro 交互设计与流程图', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

subtitle = doc.add_paragraph('AFAC2026 方向四：Agentic AI')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

sections = [
    ('一、系统总体架构', img_dir / 'diagram01-arch.png',
     '系统采用前后端分离架构，基于 Docker Compose 一键部署。前端通过 Nginx 反向代理访问后端 RESTful API 和 WebSocket。'),
    ('二、用户操作主流程', img_dir / 'diagram02-flow.png',
     '用户从登录到完成投资分析的端到端流程，涵盖输入阶段、AI 分析阶段与报告输出阶段。'),
    ('三、DAG 并行执行数据流', img_dir / 'diagram03-dag.png',
     'Orchestrator 完成意图识别后，市场分析师与情绪扫描器并行执行，风险经理与组合顾问串行消费前序结果，最终通过 WebSocket 推送报告。'),
    ('四、Agent 状态机', img_dir / 'diagram04-state.png',
     '每个 Agent 任务在生命周期内经历 PENDING → RUNNING → COMPLETED / FAILED，支持超时降级与重试。'),
]

for heading, img_path, desc in sections:
    h = doc.add_heading(heading, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    p = doc.add_paragraph(desc)
    p.runs[0].font.size = Pt(10.5)
    p.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f'[图片未找到: {img_path}]')

    doc.add_paragraph()

# Footer note
footer = doc.add_paragraph('—— FinAgent Pro · 让 AI 成为每个投资者的专业顾问 ——')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(10)
footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save(out_path)
print(f'Updated {out_path}')
